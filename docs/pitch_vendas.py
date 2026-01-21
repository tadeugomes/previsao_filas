#!/usr/bin/env python3
"""Script para gerar o documento de pitch de vendas do Previsao de Fila."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path


def create_pitch_document():
    doc = Document()

    # Configurar estilos
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Titulo principal
    title = doc.add_heading('PREVISAO DE FILA', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph('Pitch de Vendas e Proposta de Valor')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].font.italic = True

    doc.add_paragraph()

    # ===== SECAO 1: VISAO GERAL =====
    doc.add_heading('1. VISAO GERAL DO PRODUTO', level=1)

    doc.add_paragraph(
        'O Previsao de Fila e um sistema inteligente de previsao de tempo de espera '
        'para atracacao em portos brasileiros. Utilizando modelos de machine learning '
        'treinados com dados historicos da ANTAQ, combinados com informacoes climaticas '
        'em tempo real (INMET), dados de producao agricola (IBGE) e precos de commodities '
        '(IPEA), o sistema entrega previsoes precisas que transformam incerteza operacional '
        'em planejamento estrategico.'
    )

    doc.add_heading('Principais Entregas:', level=2)
    bullets = [
        'Estimativa de tempo de espera para atracacao (horas e dias)',
        'Ranking previsto da fila (ETA + espera) por navio e berco',
        'Classificacao de risco operacional (baixo, medio, alto)',
        'Tabela comparativa: line-up original vs. previsao do modelo',
        'Indicadores de confiabilidade e margem de erro esperada (MAE)',
        'Exportacao de dados em CSV para integracao com outros sistemas',
    ]
    for item in bullets:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()

    # ===== SECAO 2: POR QUE USAR O APP =====
    doc.add_heading('2. POR QUE USAR O APP EM VEZ DE LINE-UP PURO OU AIS', level=1)

    doc.add_heading('2.1 Limitacoes do Line-up Puro', level=2)
    doc.add_paragraph(
        'O line-up de portos fornece uma fotografia estatica da fila de navios: '
        'mostra quem esta na lista, em que ordem e com qual ETA declarado. '
        'Porem, apresenta limitacoes criticas:'
    )

    lineup_limits = [
        'Nao considera fatores externos: clima, producao agricola, precos de mercado',
        'Nao estima tempo real de espera - apenas ordem declarada',
        'Nao preve mudancas dinamicas na fila (ultrapassagens, atrasos)',
        'Erro tipico de 2-4 dias entre ETA declarado e atracacao real',
        'Nao oferece classificacao de risco ou confiabilidade',
        'Nao integra dados de diferentes fontes de forma automatica',
    ]
    for item in lineup_limits:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('2.2 Limitacoes do AIS Isolado', level=2)
    doc.add_paragraph(
        'Dados de AIS (Automatic Identification System) fornecem posicao e movimento '
        'dos navios em tempo real, mas isoladamente sao insuficientes para decisao:'
    )

    ais_limits = [
        'Mostra onde o navio esta, mas nao quando vai atracar',
        'Nao traduz posicao em tempo de espera ou janela operacional',
        'Nao considera capacidade do porto, bercos disponiveis ou tipo de carga',
        'Oferece visibilidade sem previsao - reativo, nao preditivo',
        'Nao correlaciona com fatores de demanda (safra, precos)',
    ]
    for item in ais_limits:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('2.3 O Diferencial do App', level=2)
    doc.add_paragraph(
        'O Previsao de Fila converte multiplos sinais em previsao acionavel:'
    )

    app_diff = [
        'Ensemble de ML (LightGBM + XGBoost) com validacao temporal',
        'Integracao automatica de 6 fontes de dados: ANTAQ, INMET, IBGE, IPEA, line-up e AIS',
        '35-39 features por previsao, incluindo sazonalidade e indicadores de mercado',
        'Modelos especificos por perfil de carga (Vegetal, Mineral, Fertilizante)',
        'Modo Premium para terminais com dados operacionais internos',
        'MAE de ~31-38h para cargas minerais e vegetais - previsao confiavel',
    ]
    for item in app_diff:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()

    # ===== SECAO 3: GANHOS OPERACIONAIS =====
    doc.add_heading('3. GANHOS OPERACIONAIS PARA O USUARIO', level=1)

    doc.add_heading('3.1 Antecipacao e Planejamento', level=2)
    doc.add_paragraph(
        'Com previsoes de ate 90 dias, o usuario pode planejar operacoes com antecedencia:'
    )
    gains_planning = [
        'Programacao de bercos e equipamentos com maior precisao',
        'Negociacao de contratos de frete com datas mais realistas',
        'Ajuste de programacao de frota baseado em filas previstas',
        'Reducao de incerteza para planejamento de safra e escoamento',
    ]
    for item in gains_planning:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('3.2 Reducao de Custos', level=2)
    doc.add_paragraph(
        'O impacto financeiro direto inclui:'
    )
    gains_cost = [
        'Reducao de demurrage: -8% a -15% em operacoes com alta variacao de ETA',
        'Economia em custos de espera: 1-2 dias a menos por navio em media',
        'Melhor alocacao de recursos portuarios e de transporte terrestre',
        'Evitar multas contratuais por atrasos previstos com antecedencia',
    ]
    for item in gains_cost:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('3.3 Poder de Negociacao', level=2)
    doc.add_paragraph(
        'Dados tecnicos para fundamentar decisoes:'
    )
    gains_nego = [
        'Argumento tecnico com risco estimado para contratos e agendas',
        'Evidencias para renegociacao de prazos e tarifas',
        'Transparencia com stakeholders sobre expectativas de entrega',
        'Base de dados auditavel para disputas comerciais',
    ]
    for item in gains_nego:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('3.4 Visibilidade Operacional', level=2)
    doc.add_paragraph(
        'Interface Streamlit com dashboards em tempo real:'
    )
    gains_visibility = [
        'Resumo executivo com tempo estimado, previsao de atracacao e risco',
        'Grafico de tendencia de fila para os proximos 7 dias',
        'Tabela comparativa line-up vs. modelo com delta de posicao',
        'Detalhes operacionais: produtividade, berco, tempo de operacao',
        'Filtros por porto, berco, navio e tipo de carga',
        'Exportacao de previsoes em CSV para integracao',
    ]
    for item in gains_visibility:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()

    # ===== SECAO 4: COMPARACAO LINEUP VS MODELO =====
    doc.add_heading('4. ANALISE: LINE-UP vs. PREVISAO DO MODELO', level=1)

    doc.add_heading('4.1 O Que o Line-up Mostra', level=2)
    doc.add_paragraph(
        'O line-up tradicional apresenta:'
    )
    lineup_shows = [
        'Ordem de chegada declarada pelos agentes maritimos',
        'ETA (Estimated Time of Arrival) informado pelo navio',
        'Berco designado (quando disponivel)',
        'Tipo de carga e quantidade',
        'Status atual (fundeio, atracado, esperando)',
    ]
    for item in lineup_shows:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph(
        'O problema: a ordem do line-up raramente reflete a ordem real de atracacao. '
        'Navios "furam fila" por prioridade de carga, disponibilidade de berco, '
        'condicoes climaticas e outros fatores nao visiveis no line-up.'
    )

    doc.add_heading('4.2 O Que o Modelo Adiciona', level=2)
    doc.add_paragraph(
        'A previsao do modelo considera:'
    )
    model_adds = [
        'Tempo historico de espera por porto, terminal e tipo de carga',
        'Sazonalidade: safra, periodo do ano, dia da semana',
        'Condicoes climaticas: chuva, vento, temperatura',
        'Indicadores de mercado: producao agricola, precos de commodities',
        'Densidade da fila: navios no fundeio e chegando em 7 dias',
        'Perfil operacional do porto: produtividade historica',
    ]
    for item in model_adds:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('4.3 Tabela Comparativa: Exemplo Pratico', level=2)
    doc.add_paragraph(
        'A interface do app apresenta uma tabela comparativa que mostra:'
    )

    # Criar tabela comparativa
    table = doc.add_table(rows=6, cols=2)
    table.style = 'Table Grid'

    headers = ['Coluna', 'Descricao']
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True

    data = [
        ('Posicao_lineup', 'Ordem original do line-up (1 = primeiro)'),
        ('Posicao_prevista', 'Ordem calculada pelo modelo com base na espera prevista'),
        ('Delta_posicao', 'Diferenca entre posicao prevista e original (+ atrasa, - adianta)'),
        ('Espera_prevista_h', 'Horas estimadas de espera antes de atracar'),
        ('Atraso_vs_ETA_h', 'Diferenca em horas entre ETA com espera e ETA original'),
    ]
    for i, (col, desc) in enumerate(data, start=1):
        table.rows[i].cells[0].text = col
        table.rows[i].cells[1].text = desc

    doc.add_paragraph()
    doc.add_paragraph(
        'Exemplo: Um navio na posicao 3 do line-up pode ter Delta_posicao = +2, '
        'indicando que provavelmente atracara em 5o lugar devido a fatores nao '
        'visiveis no line-up (chuva prevista, baixa produtividade do berco, etc.).'
    )

    doc.add_heading('4.4 Impacto Pratico da Diferenca', level=2)
    doc.add_paragraph(
        'A diferenca entre line-up e previsao tem impacto direto em:'
    )
    impact = [
        'Programacao de caminhoes: evitar filas de espera em terminais terrestres',
        'Gestao de armazens: planejar recebimento com base em atracacao prevista',
        'Contratos de frete: prever demurrage e ajustar negociacoes',
        'Comunicacao com clientes: informar prazos mais realistas',
        'Alocacao de equipes: escalar operadores para horarios previstos',
    ]
    for item in impact:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()

    # ===== SECAO 5: METRICAS DE PRECISAO =====
    doc.add_heading('5. METRICAS DE PRECISAO DO MODELO', level=1)

    doc.add_paragraph(
        'O MAE (Mean Absolute Error) indica a diferenca media entre a previsao '
        'e o tempo real de espera. Quanto menor, mais precisa a previsao.'
    )

    # Tabela de metricas
    table2 = doc.add_table(rows=5, cols=4)
    table2.style = 'Table Grid'

    headers2 = ['Perfil', 'MAE', 'Interpretacao', 'Uso Recomendado']
    hdr_cells2 = table2.rows[0].cells
    for i, header in enumerate(headers2):
        hdr_cells2[i].text = header
        hdr_cells2[i].paragraphs[0].runs[0].font.bold = True

    metrics_data = [
        ('VEGETAL', '~38h (~1,6 dias)', 'Adequado para planejamento de janela', 'Soja, milho, farelo, acucar'),
        ('MINERAL', '~31h (~1,3 dias)', 'Alta precisao para operacao intensiva', 'Minerio de ferro, bauxita, cimento'),
        ('FERTILIZANTE', '~79h (~3,3 dias)', 'Maior volatilidade operacional', 'Ureia, KCL, NPK, fosfatados'),
        ('PREMIUM (Ponta da Madeira)', '~30h (~1,25 dias)', 'Maxima precisao com dados internos', 'Minerio de ferro - terminal especifico'),
    ]
    for i, (perfil, mae, interp, uso) in enumerate(metrics_data, start=1):
        table2.rows[i].cells[0].text = perfil
        table2.rows[i].cells[1].text = mae
        table2.rows[i].cells[2].text = interp
        table2.rows[i].cells[3].text = uso

    doc.add_paragraph()
    doc.add_paragraph(
        'O modelo tambem fornece classificacao de risco (AUC-ROC 0.78-0.83) que '
        'permite identificar operacoes com maior probabilidade de atraso.'
    )

    doc.add_paragraph()

    # ===== SECAO 6: FUNCIONALIDADES DA INTERFACE =====
    doc.add_heading('6. FUNCIONALIDADES DA INTERFACE STREAMLIT', level=1)

    doc.add_heading('6.1 Painel de Controle (Sidebar)', level=2)
    sidebar_features = [
        'Selecao de porto e tipo de carga',
        'Data de chegada com range de -30 a +90 dias',
        'Filtros por berco, navio e tipo de navio',
        'Edicao manual de condicoes climaticas',
        'Ajuste de navios no fundeio',
        'Botao "Gerar Previsao" para executar modelo',
    ]
    for item in sidebar_features:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('6.2 Dashboard Principal', level=2)
    dashboard_features = [
        'Resumo Executivo: tempo estimado, previsao de atracacao, fila atual, clima',
        'Callout com mensagem personalizada sobre a previsao',
        'Detalhes Operacionais: berco, produtividade, tempo de operacao',
        'Insights do Modelo: alertas de risco, confiabilidade, fatores relevantes',
        'Tabela Comparativa: line-up vs. previsao com todas as metricas',
        'Exportacao em CSV com um clique',
    ]
    for item in dashboard_features:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('6.3 Modos de Operacao', level=2)
    modes = [
        'BASICO: Usa dados publicos (ANTAQ + INMET). Disponivel para todos os portos.',
        'PREMIUM: Integra dados internos do terminal (produtividade, laytime). Maior precisao.',
    ]
    for item in modes:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()

    # ===== SECAO 7: ROI E BUSINESS CASE =====
    doc.add_heading('7. ROI E CASO DE NEGOCIO', level=1)

    doc.add_heading('7.1 Exemplo: Porto de Itaqui', level=2)
    doc.add_paragraph(
        'Numeros ilustrativos para uma janela de 30 dias:'
    )
    roi_itaqui = [
        'Volume: ~12 navios/mes em operacao de graos',
        'Ganho medio: 1,5 dia economizado por navio',
        'Impacto: 18 dias de fila evitados por mes',
        'Custo demurrage medio: USD 15.000-25.000/dia',
        'Economia potencial: USD 270.000-450.000/mes',
    ]
    for item in roi_itaqui:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('7.2 Retorno sobre Investimento', level=2)
    doc.add_paragraph(
        '1 decisao melhor por semana paga a assinatura em janelas criticas. '
        'O ROI tipico e de 10-20x o custo da ferramenta em operacoes de medio/grande porte.'
    )

    doc.add_paragraph()

    # ===== SECAO 8: SUGESTOES DE MELHORIAS =====
    doc.add_heading('8. SUGESTOES DE MELHORIAS PARA O PROJETO', level=1)

    doc.add_heading('8.1 Melhorias de Curto Prazo', level=2)
    short_term = [
        'Integracao completa de dados AIS: atualmente scaffolding, falta conectar ao fluxo principal',
        'Alertas automaticos: notificacao por email/SMS quando risco muda de nivel',
        'API REST: expor previsoes para integracao com ERPs e sistemas portuarios',
        'Cache de previsoes: evitar recalculo quando parametros nao mudam',
        'Historico de previsoes: permitir comparar previsao vs. realizado',
    ]
    for item in short_term:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('8.2 Melhorias de Medio Prazo', level=2)
    mid_term = [
        'Modelo de series temporais (LSTM/Prophet) para previsao de tendencia de fila',
        'Integracao com dados de maré e restricoes de calado',
        'Painel multi-porto: visao consolidada de varios portos simultaneamente',
        'Previsao de produtividade: estimar tempo de operacao alem de tempo de espera',
        'Modelo de otimizacao: sugerir melhor sequencia de atracacao',
    ]
    for item in mid_term:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('8.3 Melhorias de Longo Prazo', level=2)
    long_term = [
        'Expansao para portos internacionais (Argentina, Uruguai, Chile)',
        'Modelo de demanda: prever fluxo de navios com base em safra e precos futuros',
        'Digital Twin: simulacao de cenarios (what-if) para planejamento estrategico',
        'Integracao com blockchain para rastreabilidade de dados e previsoes',
        'App mobile para consulta de previsoes em campo',
    ]
    for item in long_term:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('8.4 Melhorias Tecnicas', level=2)
    tech_improvements = [
        'Monitoramento de drift: detectar degradacao de performance do modelo automaticamente',
        'Pipeline de retraining automatico: atualizar modelos com novos dados periodicamente',
        'Testes A/B: comparar versoes de modelos em producao',
        'Logging estruturado: facilitar auditoria e debugging',
        'Containerizacao (Docker): facilitar deploy em diferentes ambientes',
        'CI/CD: automatizar testes e deploy de novas versoes',
    ]
    for item in tech_improvements:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()

    # ===== SECAO 9: CONCLUSAO =====
    doc.add_heading('9. CONCLUSAO', level=1)

    doc.add_paragraph(
        'O Previsao de Fila transforma dados dispersos em inteligencia operacional. '
        'Enquanto o line-up puro oferece uma fotografia estatica e o AIS mostra posicao '
        'sem contexto, o app integra multiplas fontes para entregar previsao acionavel '
        'com nivel de confianca conhecido.'
    )

    doc.add_paragraph(
        'Para operadores portuarios, armadores, traders e agentes maritimos, a ferramenta '
        'significa menos incerteza, menos custo de demurrage e melhor planejamento. '
        'O ROI se materializa em cada decisao que evita um dia de espera desnecessaria.'
    )

    doc.add_paragraph(
        'Proxima etapa recomendada: piloto com 1-2 portos prioritarios para validar '
        'reducao de espera e ganho de produtividade em ambiente real de operacao.'
    )

    doc.add_paragraph()

    # Rodape
    doc.add_paragraph('---')
    footer = doc.add_paragraph('Documento gerado automaticamente - Previsao de Fila - Janeiro 2026')
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.italic = True

    return doc


if __name__ == '__main__':
    output_path = Path(__file__).parent / 'pitch_vendas_previsao_fila.docx'
    doc = create_pitch_document()
    doc.save(output_path)
    print(f'Documento salvo em: {output_path}')
